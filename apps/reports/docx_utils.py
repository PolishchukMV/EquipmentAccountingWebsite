"""
Утилиты для генерации отчётов в формате Word.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils import timezone


def generate_equipment_docx(equipment_queryset):
    """
    Генерирует отчёт по оборудованию в формате Word.
    
    Args:
        equipment_queryset: QuerySet с объектами оборудования
        
    Returns:
        bytes: Буфер с данными Word файла
    """
    from io import BytesIO
    
    doc = Document()
    
    # Заголовок документа
    heading = doc.add_heading('Отчёт по оборудованию', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата формирования
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f'Дата формирования: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
    
    doc.add_paragraph()
    
    # Таблица
    table = doc.add_table(rows=1, cols=12)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    headers = [
        'Инв. номер',
        'Наименование',
        'Категория',
        'Статус',
        'Серийный номер',
        'Производитель',
        'Модель',
        'Дата покупки',
        'Цена',
        'Подразделение',
        'Ответственное лицо',
        'Местоположение'
    ]
    
    for i, header in enumerate(headers):
        run = hdr_cells[i].add_run(header)
        run.bold = True
    
    # Запись данных
    for equipment in equipment_queryset:
        row_cells = table.add_row().cells
        row_cells[0].text = equipment.inventory_number
        row_cells[1].text = equipment.name
        row_cells[2].text = equipment.category.name if equipment.category else ''
        row_cells[3].text = equipment.get_status_display()
        row_cells[4].text = equipment.serial_number or ''
        row_cells[5].text = equipment.manufacturer or ''
        row_cells[6].text = equipment.model or ''
        row_cells[7].text = equipment.purchase_date.strftime('%d.%m.%Y') if equipment.purchase_date else ''
        row_cells[8].text = str(equipment.purchase_price) if equipment.purchase_price else ''
        row_cells[9].text = equipment.department.name if equipment.department else ''
        row_cells[10].text = equipment.responsible_person.get_full_name() if equipment.responsible_person else ''
        row_cells[11].text = equipment.location or ''
    
    # Сохранение в буфер
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output