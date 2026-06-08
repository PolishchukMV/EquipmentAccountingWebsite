"""
Тесты для утилит экспорта.
"""

from django.test import TestCase
from django.contrib.auth import get_user_model
from apps.departments.models import Department
from apps.equipment.models import Category, Equipment
from apps.reports.utils import generate_equipment_excel
from apps.reports.docx_utils import generate_equipment_docx


UserModel = get_user_model()


class EquipmentExcelExportTest(TestCase):
    """Тесты экспорта в Excel."""
    
    def setUp(self):
        """Настройка тестового окружения."""
        self.department = Department.objects.create(name='IT-отдел', code='IT')
        self.category = Category.objects.create(name='Компьютеры', code='COMP')
        
        self.equipment = Equipment.objects.create(
            inventory_number='000001',
            name='Ноутбук Dell',
            category=self.category,
            status='in_use',
            serial_number='SN123',
            manufacturer='Dell',
            model='Latitude',
            purchase_price=85000,
            department=self.department,
            location='Кабинет 301'
        )
    
    def test_generate_excel(self):
        """Тест генерации Excel файла."""
        equipment_qs = Equipment.objects.all()
        
        excel_file = generate_equipment_excel(equipment_qs)
        
        self.assertIsNotNone(excel_file)
        self.assertTrue(len(excel_file.read()) > 0)
    
    def test_excel_content(self):
        """Тест содержимого Excel файла."""
        from openpyxl import load_workbook
        
        equipment_qs = Equipment.objects.all()
        excel_file = generate_equipment_excel(equipment_qs)
        
        # Загружаем workbook из буфера
        wb = load_workbook(filename=excel_file)
        ws = wb.active
        
        # Проверяем заголовки
        headers = [cell.value for cell in ws[1]]
        self.assertIn('Инвентарный номер', headers)
        self.assertIn('Наименование', headers)
        self.assertIn('Категория', headers)
        
        # Проверяем данные
        row2 = [cell.value for cell in ws[2]]
        self.assertIn('000001', row2)
        self.assertIn('Ноутбук Dell', row2)


class EquipmentWordExportTest(TestCase):
    """Тесты экспорта в Word."""
    
    def setUp(self):
        """Настройка тестового окружения."""
        self.department = Department.objects.create(name='IT-отдел', code='IT')
        self.category = Category.objects.create(name='Компьютеры', code='COMP')
        
        self.equipment = Equipment.objects.create(
            inventory_number='000001',
            name='Ноутбук Dell',
            category=self.category,
            status='in_use',
            serial_number='SN123',
            manufacturer='Dell',
            model='Latitude',
            purchase_price=85000
        )
    
    def test_generate_docx(self):
        """Тест генерации Word файла."""
        equipment_qs = Equipment.objects.all()
        
        docx_file = generate_equipment_docx(equipment_qs)
        
        self.assertIsNotNone(docx_file)
        self.assertTrue(len(docx_file.read()) > 0)
    
    def test_docx_content(self):
        """Тест содержимого Word файла."""
        from docx import Document
        
        equipment_qs = Equipment.objects.all()
        docx_file = generate_equipment_docx(equipment_qs)
        
        # Загружаем документ из буфера
        docx_file.seek(0)
        doc = Document(docx_file)
        
        # Проверяем заголовок
        headings = [p.text for p in doc.paragraphs if p.style.name == 'Heading 1']
        self.assertIn('Отчёт по оборудованию', headings)
        
        # Проверяем таблицу
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        
        # Проверяем данные в таблице
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
        self.assertIn('000001', table_text)
        self.assertIn('Ноутбук Dell', table_text)