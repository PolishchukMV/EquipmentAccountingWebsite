"""
Утилиты для экспорта данных в различные форматы.
"""

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from django.utils import timezone
from io import BytesIO


def generate_equipment_excel(equipment_queryset):
    """
    Генерирует отчёт по оборудованию в формате Excel.
    
    Args:
        equipment_queryset: QuerySet с объектами оборудования
        
    Returns:
        bytes: Буфер с данными Excel файла
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Оборудование"
    
    # Стили
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Заголовки
    headers = [
        'Инвентарный номер',
        'Наименование',
        'Категория',
        'Статус',
        'Серийный номер',
        'Производитель',
        'Модель',
        'Дата покупки',
        'Цена',
        'Подразделение',
        'Ответственное лицо',
        'Местоположение'
    ]
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    # Данные
    for row, equipment in enumerate(equipment_queryset, 2):
        ws.cell(row=row, column=1, value=equipment.inventory_number)
        ws.cell(row=row, column=2, value=equipment.name)
        ws.cell(row=row, column=3, value=equipment.category.name if equipment.category else '')
        ws.cell(row=row, column=4, value=equipment.get_status_display())
        ws.cell(row=row, column=5, value=equipment.serial_number or '')
        ws.cell(row=row, column=6, value=equipment.manufacturer or '')
        ws.cell(row=row, column=7, value=equipment.model or '')
        ws.cell(row=row, column=8, value=equipment.purchase_date.strftime('%d.%m.%Y') if equipment.purchase_date else '')
        ws.cell(row=row, column=9, value=float(equipment.purchase_price) if equipment.purchase_price else '')
        ws.cell(row=row, column=10, value=equipment.department.name if equipment.department else '')
        ws.cell(row=row, column=11, value=equipment.responsible_person.get_full_name() if equipment.responsible_person else '')
        ws.cell(row=row, column=12, value=equipment.location or '')
        
        # Применяем границы
        for col in range(1, 13):
            ws.cell(row=row, column=col).border = thin_border
    
    # Автоширина колонок
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
    
    # Сохранение в буфер
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return output


def export_equipment_to_xlsx(queryset):
    """
    Экспорт оборудования в XLSX формат (для views).
    
    Args:
        queryset: QuerySet с объектами оборудования
        
    Returns:
        HttpResponse: Ответ с файлом Excel
    """
    from django.http import HttpResponse
    
    # Обработка пустого QuerySet
    if not queryset.exists():
        from django.contrib import messages
        # Это хак для передачи сообщения - лучше использовать отдельный view
        excel_file = generate_equipment_excel(queryset)
    else:
        excel_file = generate_equipment_excel(queryset)
    
    response = HttpResponse(
        excel_file.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=equipment.xlsx'
    
    return response


def export_equipment_to_docx(equipment):
    """
    Экспорт одной единицы оборудования в DOCX формат.
    
    Args:
        equipment: Объект Equipment
        
    Returns:
        HttpResponse: Ответ с файлом Word
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.http import HttpResponse
    
    doc = Document()
    
    # Заголовок
    heading = doc.add_heading(f'Карточка оборудования №{equipment.inventory_number}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация
    doc.add_paragraph(f'Дата формирования: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('Наименование:', equipment.name),
        ('Категория:', str(equipment.category) if equipment.category else 'Не указана'),
        ('Инвентарный номер:', equipment.inventory_number),
        ('Серийный номер:', equipment.serial_number or 'Не указан'),
        ('Производитель:', equipment.manufacturer or 'Не указан'),
        ('Модель:', equipment.model or 'Не указан'),
        ('Дата покупки:', equipment.purchase_date.strftime('%d.%m.%Y') if equipment.purchase_date else 'Не указана'),
        ('Цена:', f'{equipment.purchase_price} руб.' if equipment.purchase_price else 'Не указана'),
        ('Статус:', equipment.get_status_display()),
        ('Подразделение:', str(equipment.department) if equipment.department else 'Не назначено'),
        ('Ответственный:', str(equipment.responsible_person) if equipment.responsible_person else 'Не назначен'),
        ('Местоположение:', equipment.location or 'Не указано'),
    ]
    
    for label, value in data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        row[0].paragraphs[0].runs[0].bold = True
    
    # Примечания
    if equipment.notes:
        doc.add_heading('Примечания', level=2)
        doc.add_paragraph(equipment.notes)
    
    # Сохранение
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="equipment_{equipment.inventory_number}.docx"'
    
    return response
